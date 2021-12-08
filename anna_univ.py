import docx
from docx.shared import Inches

ROWS = 62
COLS = 5
NAME_OF_FILE = "list"

doc = docx.Document()

# Creating a table object
table = doc.add_table(rows=ROWS, cols=COLS)

# looping over the images
name_list = {x: i for x, i in zip(range(202101064, 202101126), range(0, 62))}
for each in name_list:
    cell = table.rows[name_list[each]].cells[0]
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(each))

    image = table.rows[name_list[each]].cells[2]
    paragraph = image.paragraphs[0]
    run = paragraph.add_run()
    run.add_image() # work needed here

    sign = table.rows[name_list[each]].cells[3]
    paragraph = sign.paragraphs[0]
    run = paragraph.add_run()
    run.add_image()  # work needed here

    writeup = table.rows[name_list[each]].cells[4]
    paragraph = writeup.paragraphs[0]
    run = paragraph.add_run()
    run.add_image()  # work needed here

# Now save the document to a location
doc.save(f'./data/{NAME_OF_FILE}.docx')

import docx
from docx.shared import Inches

END = 346
ROWS = 9
COLS = 7
TEMPLATE = "Screenshot ({num}).png"
WIDTH = 2
HEIGHT = 3
INITIAL = 114
FINAL = 346
NAME_OF_FILE = "list"

doc = docx.Document()

# Creating a table object
table = doc.add_table(rows=ROWS * 2, cols=COLS)
  
#looping over the images
current = INITIAL
rows = 0
while True:
    if rows % 2 != 0:
        rows += 1
        continue
    columns = 0
    while True:
        try:
            print(rows, columns)
            image_cell = table.rows[rows].cells[columns]
            paragraph = image_cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(
                "./data/images/" + TEMPLATE.format(num=current), 
                width=Inches(WIDTH * 0.393701), 
                height=Inches(HEIGHT * 0.393701)
            )

            text_cell = table.rows[rows+1].cells[columns]
            paragraph = text_cell.paragraphs[0]
            paragraph.add_run(str(current))
            current += 1
        except Exception as e:
            print(e)
            print(TEMPLATE.format(num=current), 'doesn\'t exists')
            current += 1
            continue
        columns += 1
        if(columns >= COLS): break
    rows += 1
    if(rows >= ROWS * 2): break
    if(current >= FINAL): break


  
# Now save the document to a location
doc.save(f'./data/{NAME_OF_FILE}.docx')